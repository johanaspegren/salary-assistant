"""DOCX document parser — extracts body text and table cells (ORSAK/FÖRKLARING)."""

import re
from dataclasses import dataclass, field

from docx import Document


@dataclass
class TableEntry:
    header: str
    content: str
    row_index: int
    table_index: int


@dataclass
class ParsedDocument:
    filename: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[TableEntry] = field(default_factory=list)
    sections: list[dict] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        parts: list[str] = []
        for section in self.sections:
            if section["type"] == "heading":
                parts.append(f"\n## {section['text']}\n")
            elif section["type"] == "paragraph":
                parts.append(section["text"])
            elif section["type"] == "table_entry":
                parts.append(f"[{section['header']}]: {section['text']}")
        return "\n".join(parts)


def parse_docx(file_path: str, filename: str) -> ParsedDocument:
    doc = Document(file_path)
    parsed = ParsedDocument(filename=filename)

    # Extract paragraphs with heading detection
    current_heading = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name.startswith("Heading"):
            current_heading = text
            parsed.sections.append({"type": "heading", "text": text, "heading": current_heading})
        else:
            parsed.paragraphs.append(text)
            parsed.sections.append({
                "type": "paragraph",
                "text": text,
                "heading": current_heading,
            })

    # Extract tables — look for ORSAK/FÖRKLARING patterns and general Q&A
    for t_idx, table in enumerate(doc.tables):
        headers = []
        if table.rows:
            headers = [cell.text.strip().upper() for cell in table.rows[0].cells]

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0 and _is_header_row(headers):
                continue

            cells = [cell.text.strip() for cell in row.cells]

            # Try to pair cells as key-value (ORSAK->FÖRKLARING, Fråga->Svar, etc.)
            paired = _pair_cells(headers, cells, t_idx, r_idx)
            if paired:
                for entry in paired:
                    parsed.tables.append(entry)
                    parsed.sections.append({
                        "type": "table_entry",
                        "text": entry.content,
                        "header": entry.header,
                        "heading": current_heading,
                    })
            else:
                # Fallback: concatenate all non-empty cells
                combined = " | ".join(c for c in cells if c)
                if combined:
                    entry = TableEntry(
                        header="Tabell",
                        content=combined,
                        row_index=r_idx,
                        table_index=t_idx,
                    )
                    parsed.tables.append(entry)
                    parsed.sections.append({
                        "type": "table_entry",
                        "text": combined,
                        "header": "Tabell",
                        "heading": current_heading,
                    })

    return parsed


_QA_PATTERNS = re.compile(
    r"(ORSAK|FÖRKLARING|FRÅGA|SVAR|FRAGE|ANSWER|QUESTION|CAUSE|EXPLANATION)",
    re.IGNORECASE,
)


def _is_header_row(headers: list[str]) -> bool:
    return any(_QA_PATTERNS.search(h) for h in headers)


def _pair_cells(
    headers: list[str],
    cells: list[str],
    table_index: int,
    row_index: int,
) -> list[TableEntry] | None:
    if not headers or len(headers) != len(cells):
        return None

    entries = []
    for header, cell in zip(headers, cells):
        if cell:
            entries.append(
                TableEntry(
                    header=header if header else "Cell",
                    content=cell,
                    row_index=row_index,
                    table_index=table_index,
                )
            )
    return entries if entries else None
